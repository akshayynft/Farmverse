#!/usr/bin/env python3
"""
Script to convert Farmaverse Whitepaper from Markdown to DOCX format
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """Convert markdown file to DOCX format"""
    
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Set up styles
    setup_document_styles(doc)
    
    # Split content into lines
    lines = md_content.split('\n')
    
    current_paragraph = None
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_block_content:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_block_content))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_para.style = 'Quote'
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Handle headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Handle bold text
        elif '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.bold = True
        
        # Handle bullet points
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(line[line.find(' ')+1:], style='List Number')
        
        # Handle regular paragraphs
        elif line.strip():
            doc.add_paragraph(line)
        
        # Handle empty lines
        else:
            doc.add_paragraph()
    
    # Save the document
    doc.save(docx_file_path)
    print(f"Successfully converted {md_file_path} to {docx_file_path}")

def setup_document_styles(doc):
    """Set up document styles"""
    
    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    
    # Heading 1 style
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Arial'
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True
    
    # Heading 2 style
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Arial'
    heading2_style.font.size = Pt(16)
    heading2_style.font.bold = True
    
    # Heading 3 style
    heading3_style = doc.styles['Heading 3']
    heading3_style.font.name = 'Arial'
    heading3_style.font.size = Pt(14)
    heading3_style.font.bold = True
    
    # Normal paragraph style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    
    # List styles
    list_bullet_style = doc.styles['List Bullet']
    list_bullet_style.font.name = 'Arial'
    list_bullet_style.font.size = Pt(11)
    
    list_number_style = doc.styles['List Number']
    list_number_style.font.name = 'Arial'
    list_number_style.font.size = Pt(11)

def main():
    """Main function"""
    md_file = "docs/FARMAVERSE_WHITEPAPER.md"
    docx_file = "docs/FARMAVERSE_WHITEPAPER.docx"
    
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found!")
        return
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
        print(f"✅ Whitepaper successfully converted to DOCX format!")
        print(f"📄 File saved as: {docx_file}")
    except Exception as e:
        print(f"❌ Error converting file: {str(e)}")

if __name__ == "__main__":
    main() 